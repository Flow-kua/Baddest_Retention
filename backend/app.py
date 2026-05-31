from __future__ import annotations
"""
Baddest_Retention 后端服务 v3
全格式支持：图片/文档/PDF → 提取文字 → LLM 风格整理 → 存库 → 返回

启动：conda activate Baddest_Retention → python app.py
端口：5000
"""

import os
import sys
import base64
import json
import re
import time
import threading
from io import BytesIO

# 关掉 stdout 缓冲，确保日志实时显示
sys.stdout.reconfigure(line_buffering=True)

from flask import Flask, request, jsonify
from flask_cors import CORS
from PIL import Image

# 本地模块
sys.path.insert(0, os.path.dirname(__file__))
from parsers import get_parser
from database import init_db, insert_note, get_all_notes, get_note_by_id, get_note_count, update_note_file_path, update_note_meta, delete_note, get_recent_extra_dimensions
from database import DB_PATH, DATA_DIR
from profile_manager import ProfileManager
from organizer import Organizer

# ================================================================
# 全局初始化
# ================================================================
IMAGE_SAVE_DIR = os.path.join(DATA_DIR, "images")
KNOWLEDGE_DIR = os.path.join(DATA_DIR, "knowledge")
os.makedirs(IMAGE_SAVE_DIR, exist_ok=True)
os.makedirs(KNOWLEDGE_DIR, exist_ok=True)

init_db()
pm = ProfileManager(DB_PATH)

app = Flask(__name__)
CORS(app)

# ---- 请求日志 ----
@app.before_request
def log_request():
    from datetime import datetime
    ts = datetime.now().strftime("%H:%M:%S")
    print(f"[{ts}] {request.remote_addr}  {request.method}  {request.path}", flush=True)

@app.after_request
def log_response(response):
    from datetime import datetime
    ts = datetime.now().strftime("%H:%M:%S")
    print(f"[{ts}] -> {response.status_code}  {response.content_length or 0} bytes", flush=True)
    return response

# ================================================================
# 文件类型分类
# ================================================================
IMAGE_EXTS = {"png", "jpg", "jpeg", "gif", "bmp", "webp", "tiff", "tif", "ico", "heic", "heif"}
DOC_EXTS = {"md", "txt", "markdown", "rst", "log", "csv", "json", "xml", "yaml", "yml", "toml", "ini", "conf"}
DOCX_EXTS = {"docx", "doc"}
PDF_EXTS = {"pdf"}
CODE_EXTS = {"py", "js", "ts", "html", "css", "java", "c", "cpp", "go", "rs", "rb", "php", "sh", "bat", "sql", "r"}

def classify_file(filename: str) -> tuple[str, str]:
    """
    根据扩展名分类文件。
    返回 (category, source_type)
      category: "image" | "doc" | "docx" | "pdf" | "code" | "unknown"
      source_type: 用于数据库记录
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in IMAGE_EXTS:
        return "image", "photo"
    if ext in DOCX_EXTS:
        return "docx", "document"
    if ext in PDF_EXTS:
        return "pdf", "pdf"
    if ext in DOC_EXTS:
        return "doc", "text"
    if ext in CODE_EXTS:
        return "code", "code"
    return "unknown", "unknown"

# ================================================================
# 文本提取：各格式的本地提取
# ================================================================
def extract_text_from_doc(file_bytes: bytes, filename: str) -> str:
    """从纯文本文件提取内容（md, txt, csv, json 等）"""
    import chardet
    detected = chardet.detect(file_bytes)
    encoding = detected.get("encoding", "utf-8") or "utf-8"
    try:
        return file_bytes.decode(encoding, errors="ignore")
    except Exception:
        return file_bytes.decode("utf-8", errors="ignore")

def extract_text_from_docx(file_bytes: bytes, filename: str) -> str:
    """从 docx/doc 提取文本"""
    from docx import Document
    doc = Document(BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # 也提取表格内容
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            tables_text.append(" | ".join(row_text))
    result = "\n\n".join(paragraphs)
    if tables_text:
        result += "\n\n--- 表格内容 ---\n" + "\n".join(tables_text)
    return result

def extract_text_from_pdf(file_bytes: bytes, filename: str) -> str:
    """从 PDF 提取文本"""
    import fitz  # pymupdf
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_text = []
    for page in doc:
        pages_text.append(page.get_text())
    doc.close()
    return "\n\n".join(pages_text)

def extract_text_from_file(file_bytes: bytes, filename: str) -> str | None:
    """
    根据文件类型提取文本内容。
    返回 None 表示无法提取（如图片，需要走 VL 管线）。
    """
    category, _ = classify_file(filename)
    try:
        if category == "doc":
            return extract_text_from_doc(file_bytes, filename)
        elif category == "docx":
            # .doc 是旧格式，python-docx 只支持 .docx
            ext = filename.rsplit(".", 1)[-1].lower()
            if ext == "doc":
                return "[旧版 .doc 格式，暂不支持提取，请转换为 .docx]"
            return extract_text_from_docx(file_bytes, filename)
        elif category == "pdf":
            text = extract_text_from_pdf(file_bytes, filename)
            if not text.strip():
                return "[PDF 为扫描版，无可提取文字，将走图片识别]"
            return text
        elif category == "code":
            return extract_text_from_doc(file_bytes, filename)
        elif category == "image":
            return None  # 图片走 VL 管线
        else:
            return None
    except Exception as e:
        print(f"[提取失败] {filename}: {e}", flush=True)
        return None

# ================================================================
# 图片预处理
# ================================================================
def preprocess_image(file_bytes: bytes) -> tuple[str, str]:
    """图片预处理：压缩 + 转 base64"""
    img = Image.open(BytesIO(file_bytes))

    # 等比缩放到最长边 1200px
    max_size = 1200
    w, h = img.size
    if max(w, h) > max_size:
        ratio = max_size / max(w, h)
        img = img.resize((int(w * ratio), int(h * ratio)), Image.Resampling.LANCZOS)

    if img.mode == "RGBA":
        background = Image.new("RGB", img.size, (255, 255, 255))
        background.paste(img, mask=img.split()[3])
        img = background

    buf = BytesIO()
    img.save(buf, format="JPEG", quality=85)
    base64_str = base64.b64encode(buf.getvalue()).decode("utf-8")
    return base64_str, "image/jpeg"

# ================================================================
# 自动风格画像更新（后台线程）
# ================================================================
def _maybe_update_profile():
    """检查是否需要更新风格画像，如果需要则在后台线程中执行。"""
    if not pm.should_analyze():
        return
    # 不阻塞当前请求，后台线程执行
    def _do_update():
        try:
            print("[画像] 笔记数量已达标，开始分析用户风格...", flush=True)
            client = pm.build_client()
            if client is None:
                return
            notes = pm.get_recent_notes(limit=10)
            new_profile = client.analyze_user_style(notes)
            # 把当前笔记总数写入 sample_count，方便下次判断是否需要再更新
            note_count = get_note_count()
            new_profile["sample_count"] = note_count
            pm.update_profile(new_profile)
            print(f"[画像] 风格画像已更新: style={new_profile.get('style_preference')}, "
                  f"dimensions={len(new_profile.get('dimension_hints', []))}, "
                  f"habits={new_profile.get('unique_habits', [])}", flush=True)
        except Exception as e:
            print(f"[画像] 风格分析失败: {e}", flush=True)

    t = threading.Thread(target=_do_update, daemon=True)
    t.start()


# ================================================================
# 知识文件保存
# ================================================================
def _save_knowledge_file(note_id: str, organized_text: str, tags: list,
                         source_filename: str, source_type: str,
                         confidence: float, created_at: str) -> str:
    """把整理结果保存为知识目录下的 .md 文件，返回文件路径。"""
    safe_name = source_filename or "untitled"
    filename = f"{note_id}.md"
    file_path = os.path.join(KNOWLEDGE_DIR, filename)
    tag_str = ", ".join(tags) if tags else ""
    lines = [
        "---",
        f"note_id: {note_id}",
        f"source: {safe_name}",
        f"type: {source_type}",
        f"tags: [{tag_str}]",
        f"confidence: {confidence:.2f}",
        f"created_at: {created_at}",
        "---",
        "",
        organized_text,
    ]
    with open(file_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"[文件] 已保存: {file_path}", flush=True)
    return file_path


# ================================================================
# LLM 整理：统一入口
# ================================================================
def organize_with_llm(raw_text: str, source_type: str, source_filename: str, profile: dict) -> dict | None:
    """
    把提取到的文字走 LLM 管线，整理结果 → 存库 + 存为 .md 文件。
    返回 None 表示 LLM 未配置。
    """
    client = pm.build_client()
    if client is None:
        return None

    org = Organizer(client)
    try:
        result = org.process_text(raw_text, profile)
    except Exception as e:
        print(f"[LLM错误] 整理失败，保存原始文本为 fallback: {e}", flush=True)
        note_id = insert_note(
            raw_text=raw_text,
            organized_text=raw_text,
            source_type=source_type,
            source_filename=source_filename,
            structured_fields={},
            tags=["pending"],
            confidence=0.0,
            style_markers=[],
            extra_dimensions=[],
            file_path="",
        )
        from datetime import datetime
        created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        file_path = _save_knowledge_file(
            note_id, raw_text, ["pending"], source_filename, source_type, 0.0, created_at
        )
        update_note_file_path(note_id, file_path)
        return {
            "note_id": note_id,
            "file_path": file_path,
            "filename": os.path.basename(file_path),
            "raw_text": raw_text,
            "organized_text": raw_text,
            "structured_fields": {},
            "tags": ["pending"],
            "confidence": 0.0,
            "style_markers": [],
            "extra_dimensions": [],
            "evolution_suggestions": [],
        }

    note_id = insert_note(
        raw_text=raw_text,
        organized_text=result.get("organized_text", ""),
        source_type=source_type,
        source_filename=source_filename,
        structured_fields=result.get("structured_fields", {}),
        tags=result.get("tags", []),
        confidence=result.get("confidence", 0.0),
        style_markers=result.get("style_markers", []),
        extra_dimensions=result.get("extra_dimensions", []),
        file_path="",
    )

    # 保存为 .md 知识文件
    from datetime import datetime
    created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    file_path = _save_knowledge_file(
        note_id, result.get("organized_text", ""),
        result.get("tags", []), source_filename, source_type,
        result.get("confidence", 0.0), created_at
    )
    update_note_file_path(note_id, file_path)

    # 模板进化检测
    extra_dims = result.get("extra_dimensions", [])
    recent_dims = get_recent_extra_dimensions(limit=10)
    evolution = org.check_evolution(extra_dims, recent_dims) if extra_dims else []

    return {
        "note_id": note_id,
        "file_path": file_path,
        "filename": os.path.basename(file_path),
        "raw_text": raw_text,
        "organized_text": result.get("organized_text", ""),
        "structured_fields": result.get("structured_fields", {}),
        "tags": result.get("tags", []),
        "confidence": result.get("confidence", 0.0),
        "style_markers": result.get("style_markers", []),
        "extra_dimensions": extra_dims,
        "evolution_suggestions": evolution,
    }


def organize_image_with_llm(image_b64: str, mime_type: str, source_type: str, source_filename: str, profile: dict) -> dict | None:
    """图片 → VL 提取 → 整理 → 存库 + 存为 .md 文件"""
    client = pm.build_client()
    if client is None:
        return None

    org = Organizer(client)
    try:
        result = org.process_image(image_b64, profile)
    except Exception as e:
        print(f"[LLM错误] 图片处理失败: {e}", flush=True)
        from datetime import datetime
        created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        fail_msg = f"[处理失败] 文件 {source_filename} 的 AI 处理失败，请稍后手动重新处理。\n错误: {str(e)[:200]}"
        note_id = insert_note(
            raw_text=fail_msg,
            organized_text=fail_msg,
            source_type=source_type,
            source_filename=source_filename,
            structured_fields={},
            tags=["pending"],
            confidence=0.0,
            style_markers=[],
            extra_dimensions=[],
            file_path="",
        )
        file_path = _save_knowledge_file(
            note_id, fail_msg, ["pending"], source_filename, source_type, 0.0, created_at
        )
        update_note_file_path(note_id, file_path)
        return {
            "note_id": note_id,
            "file_path": file_path,
            "filename": os.path.basename(file_path),
            "raw_text": "",
            "organized_text": fail_msg,
            "structured_fields": {},
            "tags": ["pending"],
            "confidence": 0.0,
            "style_markers": [],
            "extra_dimensions": [],
            "evolution_suggestions": [],
        }

    note_id = insert_note(
        raw_text=result.get("raw_text", ""),
        organized_text=result.get("organized_text", ""),
        source_type=source_type,
        source_filename=source_filename,
        structured_fields=result.get("structured_fields", {}),
        tags=result.get("tags", []),
        confidence=result.get("confidence", 0.0),
        style_markers=result.get("style_markers", []),
        extra_dimensions=result.get("extra_dimensions", []),
        file_path="",
    )

    from datetime import datetime
    created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    file_path = _save_knowledge_file(
        note_id, result.get("organized_text", ""),
        result.get("tags", []), source_filename, source_type,
        result.get("confidence", 0.0), created_at
    )
    update_note_file_path(note_id, file_path)

    # 模板进化检测
    extra_dims = result.get("extra_dimensions", [])
    recent_dims = get_recent_extra_dimensions(limit=10)
    evolution = org.check_evolution(extra_dims, recent_dims) if extra_dims else []

    return {
        "note_id": note_id,
        "file_path": file_path,
        "filename": os.path.basename(file_path),
        "raw_text": result.get("raw_text", ""),
        "organized_text": result.get("organized_text", ""),
        "structured_fields": result.get("structured_fields", {}),
        "tags": result.get("tags", []),
        "confidence": result.get("confidence", 0.0),
        "style_markers": result.get("style_markers", []),
        "extra_dimensions": extra_dims,
        "evolution_suggestions": evolution,
    }

# ================================================================
# API 接口
# ================================================================

@app.route('/api/health', methods=['GET'])
def health_check():
    try:
        llm_ready = pm.is_configured()
        count = get_note_count()
        return jsonify({
            "status": "ok",
            "service": "baddest_retention-backend",
            "llm_ready": llm_ready,
            "note_count": count,
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"status": "error", "error": str(e)}), 500


@app.route('/api/parse', methods=['POST'])
def parse_file():
    """
    接收任意文件并解析 + AI 整理。

    图片：走 VL 管线（多模态识别 → 整理）
    文档/PDF/代码：走本地文本提取 → 整理
    """
    if 'file' not in request.files:
        return jsonify({"status": "error", "message": "未收到文件"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"status": "error", "message": "文件名为空"}), 400

    filename = file.filename
    file_bytes = file.read()
    size_kb = len(file_bytes) / 1024

    try:
        category, source_type = classify_file(filename)

        # 基础文件解析（保留前端队友的逻辑）
        parser = get_parser(filename)
        parsed = parser(file_bytes, filename)

        ai_organized = None
        profile = pm.get_profile()

        # ---- 图片：走 VL 管线 ----
        if category == "image" and pm.build_client() is not None:
            print(f"[管线] 图片 → VL识别 → 整理", flush=True)
            image_b64, mime_type = preprocess_image(file_bytes)
            ai_organized = organize_image_with_llm(
                image_b64, mime_type, source_type, filename, profile
            )

        # ---- 文档/PDF/代码：先本地提取文字，再走整理管线 ----
        elif category in ("doc", "docx", "pdf", "code") and pm.build_client() is not None:
            extracted = extract_text_from_file(file_bytes, filename)

            # PDF 扫描版无文字 → fallback 到图片识别
            if extracted and extracted.startswith("[PDF 为扫描版"):
                if pm.build_client() is not None and pm.build_client()._vl_client:
                    print(f"[管线] 扫描版PDF → VL图片识别 → 整理", flush=True)
                    image_b64, mime_type = preprocess_image(file_bytes)
                    ai_organized = organize_image_with_llm(
                        image_b64, mime_type, "pdf_scan", filename, profile
                    )
                else:
                    ai_organized = None
            elif extracted and extracted.strip():
                print(f"[管线] {category} → 本地提取 → 整理", flush=True)
                ai_organized = organize_with_llm(
                    extracted, source_type, filename, profile
                )
            else:
                ai_organized = None

        # 整理成功后，后台检查是否需要更新风格画像
        if ai_organized is not None:
            _maybe_update_profile()

        return jsonify({
            "status": "success",
            "filename": filename,
            "category": category,
            "file_type": parsed.get("type", "unknown"),
            "size_kb": round(size_kb, 2),
            "parsed_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            "parsed": parsed,
            "ai_organized": ai_organized,
        })

    except Exception as e:
        print(f"[错误] {filename}: {e}", flush=True)
        return jsonify({
            "status": "error",
            "message": f"解析失败: {str(e)}",
            "filename": filename,
        }), 500


@app.route('/api/text', methods=['POST'])
def process_text():
    """接收纯文本 → LLM 整理 → 存库 → 返回"""
    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({"status": "error", "message": "请提供 text 字段"}), 400

    text = data['text'].strip()
    if not text:
        return jsonify({"status": "error", "message": "文本内容为空"}), 400

    client = pm.build_client()
    if client is None:
        return jsonify({"status": "error", "message": "LLM 未配置，请先调用 /api/setup-llm"}), 400

    try:
        ai_organized = organize_with_llm(text, "text", "", pm.get_profile())
        if ai_organized is None:
            return jsonify({"status": "error", "message": "整理失败"}), 500

        # 整理成功后，后台检查是否需要更新风格画像
        _maybe_update_profile()

        return jsonify({
            "status": "success",
            "parsed_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            **ai_organized,
        })
    except Exception as e:
        return jsonify({"status": "error", "message": f"整理失败: {str(e)}"}), 500


@app.route('/api/notes', methods=['GET'])
def list_notes():
    notes = get_all_notes()
    return jsonify({"status": "success", "notes": notes, "count": len(notes)})


@app.route('/api/notes/<note_id>', methods=['GET'])
def get_note(note_id):
    note = get_note_by_id(note_id)
    if not note:
        return jsonify({"status": "error", "message": "笔记不存在"}), 404
    return jsonify({"status": "success", "note": note})


@app.route('/api/profile', methods=['GET'])
def get_profile():
    profile = pm.get_profile()
    llm_config = pm.get_llm_config()
    safe_config = {
        "base_url": llm_config.get("base_url", ""),
        "vl_model": llm_config.get("vl_model", ""),
        "vl_base_url": llm_config.get("vl_base_url", "") or llm_config.get("base_url", ""),
        "vl_configured": bool(llm_config.get("vl_api_key") or llm_config.get("api_key")),
        "text_model": llm_config.get("text_model", ""),
        "text_base_url": llm_config.get("text_base_url", "") or llm_config.get("base_url", ""),
        "text_configured": bool(llm_config.get("text_api_key") or llm_config.get("api_key")),
        "configured": bool(llm_config.get("text_api_key") or llm_config.get("api_key")),
    }
    return jsonify({"status": "success", "profile": profile, "llm_config": safe_config})


@app.route('/api/setup-llm', methods=['POST'])
def setup_llm():
    data = request.get_json()
    try:
        pm.set_llm_config(
            api_key=data.get('api_key', '').strip(),
            base_url=data.get('base_url', '').strip(),
            vl_model=data.get('vl_model', '').strip(),
            text_model=data.get('text_model', '').strip(),
            vl_api_key=data.get('vl_api_key', '').strip(),
            vl_base_url=data.get('vl_base_url', '').strip(),
            text_api_key=data.get('text_api_key', '').strip(),
            text_base_url=data.get('text_base_url', '').strip(),
        )
        client = pm.build_client()
        if client is None:
            return jsonify({"status": "error", "message": "配置不完整"}), 400
        vl_ok = bool(client._vl_client and client.vl_model)
        return jsonify({"status": "success", "message": "LLM 配置成功", "vl_ready": vl_ok, "text_ready": True})
    except Exception as e:
        return jsonify({"status": "error", "message": f"配置失败: {str(e)}"}), 400


@app.route('/api/profile/update', methods=['POST'])
def update_profile():
    note_count = get_note_count()
    if note_count < 5:
        return jsonify({"status": "error", "message": f"笔记数量不足（当前 {note_count} 条，需要至少 5 条）"}), 400
    client = pm.build_client()
    if client is None:
        return jsonify({"status": "error", "message": "LLM 未配置"}), 400
    try:
        notes = pm.get_recent_notes(limit=10)
        new_profile = client.analyze_user_style(notes)
        pm.update_profile(new_profile)
        return jsonify({"status": "success", "message": "风格画像已更新", "new_profile": pm.get_profile()})
    except Exception as e:
        return jsonify({"status": "error", "message": f"分析失败: {str(e)}"}), 500


# ================================================================
# 知识文件管理 API
# ================================================================

@app.route('/api/files', methods=['GET'])
def list_files():
    """列出知识目录下所有 .md 文件及其元数据。"""
    files = []
    notes = get_all_notes()
    for note in notes:
        note_id = note["id"]
        file_path = os.path.join(KNOWLEDGE_DIR, f"{note_id}.md")
        if os.path.exists(file_path):
            files.append({
                "note_id": note_id,
                "filename": f"{note_id}.md",
                "source_type": note.get("source_type", ""),
                "tags": note.get("tags", []),
                "confidence": note.get("confidence", 0),
                "created_at": note.get("created_at", ""),
                "summary": note.get("summary", ""),
            })
    return jsonify({"status": "success", "files": files, "count": len(files)})


@app.route('/api/files/<note_id>', methods=['GET'])
def read_file(note_id):
    """读取指定知识文件的内容 + 元数据。"""
    note = get_note_by_id(note_id)
    if not note:
        return jsonify({"status": "error", "message": "笔记不存在"}), 404

    file_path = note.get("file_path", "") or os.path.join(KNOWLEDGE_DIR, f"{note_id}.md")
    content = ""
    if os.path.exists(file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
    else:
        # 兼容旧数据：没有文件的用 organized_text
        content = note.get("organized_text", "")

    return jsonify({
        "status": "success",
        "note_id": note_id,
        "filename": os.path.basename(file_path),
        "content": content,
        "tags": note.get("tags", []),
        "source_type": note.get("source_type", ""),
        "source_filename": note.get("source_filename", ""),
        "confidence": note.get("confidence", 0),
        "style_markers": note.get("style_markers", []),
        "extra_dimensions": note.get("extra_dimensions", []),
        "created_at": note.get("created_at", ""),
    })


@app.route('/api/files/<note_id>', methods=['PUT'])
def update_file(note_id):
    """保存对知识文件的编辑。接收 JSON: {content?, title?, tags?}，至少提供一个。"""
    note = get_note_by_id(note_id)
    if not note:
        return jsonify({"status": "error", "message": "笔记不存在"}), 404

    data = request.get_json()
    if data is None:
        return jsonify({"status": "error", "message": "请提供 JSON 数据"}), 400

    new_title = data.get("title")
    new_tags = data.get("tags")
    new_content = data.get("content")

    if new_title is None and new_tags is None and new_content is None:
        return jsonify({"status": "error", "message": "请提供 content、title 或 tags 字段"}), 400

    # ---- 更新标题/标签 ----
    if new_title is not None or new_tags is not None:
        update_note_meta(
            note_id,
            source_filename=new_title.strip() if new_title else None,
            tags=new_tags if new_tags is not None else None,
        )

    # ---- 更新文件内容 ----
    file_path = note.get("file_path", "") or os.path.join(KNOWLEDGE_DIR, f"{note_id}.md")

    if new_content is not None:
        # 保留 frontmatter，只替换正文
        if os.path.exists(file_path):
            with open(file_path, "r", encoding="utf-8") as f:
                old = f.read()
            parts = old.split("---", 2)
            if len(parts) >= 3:
                new_file = parts[0] + "---" + parts[1] + "---\n\n" + new_content
            else:
                new_file = new_content
        else:
            new_file = new_content

        with open(file_path, "w", encoding="utf-8") as f:
            f.write(new_file)

        import sqlite3
        conn = sqlite3.connect(DB_PATH)
        conn.execute("UPDATE notes SET organized_text = ?, file_path = ? WHERE id = ?",
                     (new_content, file_path, note_id))
        conn.commit()
        conn.close()
        print(f"[文件] 内容已更新: {file_path}", flush=True)

    # ---- 如果标题/标签变了，重写 frontmatter ----
    if new_title is not None or new_tags is not None:
        updated_note = get_note_by_id(note_id)
        title = updated_note["source_filename"] or "untitled"
        tags = updated_note["tags"] or []
        tag_str = ", ".join(tags)
        # 读取当前正文
        body = new_content if new_content is not None else ""
        if not body and os.path.exists(file_path):
            with open(file_path, "r", encoding="utf-8") as f:
                old = f.read()
            parts = old.split("---", 2)
            body = parts[2].strip() if len(parts) >= 3 else old

        frontmatter = [
            "---",
            f"note_id: {note_id}",
            f"source: {title}",
            f"type: {updated_note['source_type']}",
            f"tags: [{tag_str}]",
            f"confidence: {updated_note['confidence']:.2f}",
            f"created_at: {updated_note['created_at']}",
            "---",
            "",
            body,
        ]
        with open(file_path, "w", encoding="utf-8") as f:
            f.write("\n".join(frontmatter))
        print(f"[文件] 元数据已更新: {file_path}", flush=True)

    return jsonify({
        "status": "success",
        "message": "文件已保存",
        "note_id": note_id,
        "file_path": file_path,
        "source_filename": (new_title.strip() if new_title else note.get("source_filename", "")),
        "tags": new_tags if new_tags is not None else note.get("tags", []),
    })


@app.route('/api/files', methods=['POST'])
def create_file():
    """
    手动创建新知识文件。
    接收 JSON: {title: "标题", content: "可选内容", tags: [...], source_type: "manual"}
    """
    data = request.get_json()
    if not data:
        return jsonify({"status": "error", "message": "请提供文件信息"}), 400

    title = data.get("title", "未命名").strip()
    content = data.get("content", "").strip()
    tags = data.get("tags", [])
    source_type = data.get("source_type", "manual")

    from datetime import datetime
    note_id = insert_note(
        raw_text=content,
        organized_text=content,
        source_type=source_type,
        source_filename=title,
        tags=tags,
        file_path="",
    )
    created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    file_path = _save_knowledge_file(
        note_id, content, tags, title, source_type, 1.0, created_at
    )
    update_note_file_path(note_id, file_path)

    return jsonify({
        "status": "success",
        "message": "文件已创建",
        "note_id": note_id,
        "filename": os.path.basename(file_path),
        "file_path": file_path,
    })


@app.route('/api/files/<note_id>', methods=['DELETE'])
def delete_file(note_id):
    """删除知识文件及其数据库记录。"""
    note = get_note_by_id(note_id)
    if not note:
        return jsonify({"status": "error", "message": "笔记不存在"}), 404

    file_path = note.get("file_path", "") or os.path.join(KNOWLEDGE_DIR, f"{note_id}.md")
    if os.path.exists(file_path):
        os.remove(file_path)
        print(f"[文件] 已删除: {file_path}", flush=True)

    delete_note(note_id)
    return jsonify({"status": "success", "message": "文件已删除", "note_id": note_id})


if __name__ == '__main__':
    print("=" * 50)
    print("  Baddest_Retention 后端服务 v4")
    print("  地址: http://localhost:5000")
    print(f"  知识目录: {KNOWLEDGE_DIR}")
    print("  支持格式:")
    print("    图片: png jpg jpeg gif bmp webp tiff svg heic")
    print("    文档: md txt docx doc pdf csv json yaml")
    print("    代码: py js ts html css java c cpp go rs")
    print("=" * 50)
    app.run(host='0.0.0.0', port=5000, debug=False)